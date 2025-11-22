"""
Extração Híbrida de Texto
Suporta PDF, Word (.docx), e fallback para Textract
Economia estimada: ~80% em custos de extração
"""

import PyPDF2
import mimetypes
from docx import Document
from typing import Optional, Tuple
from src.config import (
    ENABLE_PYPDF2,
    ENABLE_TEXTRACT,
    PYPDF2_FALLBACK_TO_TEXTRACT,
    MIN_EXTRACTED_TEXT_LENGTH
)


def extract_text_with_pypdf2(pdf_path: str) -> Optional[str]:
    """
    Extrai texto com PyPDF2 (grátis, rápido)
    
    Args:
        pdf_path: Caminho do arquivo PDF
        
    Returns:
        str: Texto extraído ou None se falhar
    """
    if not ENABLE_PYPDF2:
        return None
    
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            extracted_text = '\n'.join(text_parts)
            
            if len(extracted_text) >= MIN_EXTRACTED_TEXT_LENGTH:
                print(f'✅ PyPDF2: {len(extracted_text)} caracteres extraídos')
                return extracted_text
            else:
                print(f'⚠️ PyPDF2: Texto muito curto ({len(extracted_text)} chars)')
                return None
                
    except Exception as e:
        print(f'❌ PyPDF2 falhou: {e}')
        return None


def extract_text_with_textract(textract_client, s3_bucket: str, s3_key: str) -> Optional[str]:
    """
    Extrai texto com AWS Textract (pago, mas preciso)
    
    Args:
        textract_client: Cliente boto3 Textract
        s3_bucket: Bucket do S3
        s3_key: Chave do objeto
        
    Returns:
        str: Texto extraído ou None se falhar
    """
    if not ENABLE_TEXTRACT:
        return None
    
    try:
        response = textract_client.detect_document_text(
            Document={
                'S3Object': {
                    'Bucket': s3_bucket,
                    'Name': s3_key
                }
            }
        )
        
        # Extrair texto dos blocos
        text_parts = []
        for block in response.get('Blocks', []):
            if block['BlockType'] == 'LINE':
                text_parts.append(block.get('Text', ''))
        
        extracted_text = '\n'.join(text_parts)
        
        print(f'✅ Textract: {len(extracted_text)} caracteres extraídos')
        return extracted_text
        
    except Exception as e:
        print(f'❌ Textract falhou: {e}')
        return None


def extract_text_from_docx(docx_path: str) -> Optional[str]:
    """
    Extrai texto de arquivos Word (.docx) usando python-docx
    
    Args:
        docx_path: Caminho do arquivo .docx
        
    Returns:
        str: Texto extraído ou None se falhar
    """
    try:
        doc = Document(docx_path)
        text_parts = []
        
        # Extrair parágrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extrair tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        extracted_text = '\n'.join(text_parts)
        
        if len(extracted_text) >= MIN_EXTRACTED_TEXT_LENGTH:
            print(f'✅ python-docx: {len(extracted_text)} caracteres extraídos')
            return extracted_text
        else:
            print(f'⚠️ python-docx: Texto muito curto ({len(extracted_text)} chars)')
            return None
            
    except Exception as e:
        print(f'❌ python-docx falhou: {e}')
        return None


def extract_text_hybrid(pdf_path: str, textract_client, s3_bucket: str, s3_key: str) -> Tuple[Optional[str], str]:
    """
    Estratégia híbrida: tenta PyPDF2 primeiro, fallback para Textract
    
    Args:
        pdf_path: Caminho local do PDF
        textract_client: Cliente boto3 Textract
        s3_bucket: Bucket S3
        s3_key: Chave S3
        
    Returns:
        Tuple[texto_extraido, metodo_usado]
    """
    # Tentativa 1: PyPDF2 (grátis)
    text = extract_text_with_pypdf2(pdf_path)
    if text:
        return text, 'pypdf2'
    
    # Tentativa 2: Textract (pago)
    if PYPDF2_FALLBACK_TO_TEXTRACT:
        print('🔄 Fazendo fallback para Textract...')
        text = extract_text_with_textract(textract_client, s3_bucket, s3_key)
        if text:
            return text, 'textract'
    
    print('❌ Falha em ambos os métodos de extração')
    return None, 'none'


def extract_text_universal(file_path: str, textract_client, s3_bucket: str, s3_key: str) -> Tuple[Optional[str], str]:
    """
    Extração universal que detecta o tipo de arquivo e usa o método apropriado
    Suporta: PDF, Word (.docx, .doc), imagens (JPG, PNG, HEIC, etc.)
    
    Args:
        file_path: Caminho local do arquivo
        textract_client: Cliente boto3 Textract
        s3_bucket: Bucket S3
        s3_key: Chave S3
        
    Returns:
        Tuple[texto_extraido, metodo_usado]
    """
    try:
        # Detectar tipo MIME pela extensão
        mime_type, _ = mimetypes.guess_type(file_path)
        
        # Fallback: detectar pela extensão manualmente
        if not mime_type:
            ext = file_path.lower().split('.')[-1]
            mime_map = {
                'pdf': 'application/pdf',
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'doc': 'application/msword',
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'png': 'image/png',
                'heic': 'image/heic',
                'heif': 'image/heif',
                'tiff': 'image/tiff',
                'tif': 'image/tiff',
                'webp': 'image/webp'
            }
            mime_type = mime_map.get(ext, 'application/octet-stream')
        
        print(f'🔍 Tipo MIME detectado: {mime_type}')
        
        # PDF: usar estratégia híbrida PyPDF2 → Textract
        if mime_type == 'application/pdf':
            return extract_text_hybrid(file_path, textract_client, s3_bucket, s3_key)
        
        # Word (.docx): python-docx → Textract
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            print('📄 Processando arquivo Word (.docx)')
            text = extract_text_from_docx(file_path)
            if text:
                return text, 'python-docx'
            
            # Fallback para Textract
            if ENABLE_TEXTRACT and PYPDF2_FALLBACK_TO_TEXTRACT:
                print('🔄 Fallback para Textract (Word)...')
                text = extract_text_with_textract(textract_client, s3_bucket, s3_key)
                if text:
                    return text, 'textract'
        
        # Word antigo (.doc): apenas Textract
        elif mime_type == 'application/msword':
            print('📄 Formato .doc antigo detectado - usando Textract')
            if ENABLE_TEXTRACT:
                text = extract_text_with_textract(textract_client, s3_bucket, s3_key)
                if text:
                    return text, 'textract'
        
        # 🆕 IMAGENS: Textract primeiro (mais barato para imagens simples)
        elif mime_type and mime_type.startswith('image/'):
            print(f'🖼️ Imagem detectada: {mime_type}')
            
            # Tentativa 1: Textract (mais barato para imagens simples)
            if ENABLE_TEXTRACT:
                print('📸 Tentando Textract para imagem...')
                text = extract_text_with_textract(textract_client, s3_bucket, s3_key)
                if text and len(text) >= MIN_EXTRACTED_TEXT_LENGTH:
                    print(f'✅ Textract extraiu {len(text)} caracteres da imagem')
                    return text, 'textract'
                else:
                    print('⚠️ Textract retornou texto insuficiente - fallback para Vision')
            
            # Tentativa 2: Vision API (mais preciso, mas mais caro)
            # Será tratado no fallback do lambda_function.py
            print('📸 Imagem será processada com Vision API (fallback)')
            return None, 'image_needs_vision'
        
        # Formato não suportado
        print(f'❌ Formato não suportado: {mime_type}')
        return None, 'unsupported'
        
    except Exception as e:
        print(f'❌ Erro na detecção de formato: {e}')
        # Fallback: tentar como PDF
        return extract_text_hybrid(file_path, textract_client, s3_bucket, s3_key)
